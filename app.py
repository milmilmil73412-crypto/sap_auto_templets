"""SAP 입찰기안문 자동생성 에이전트 — Streamlit 앱"""

import os
import sys
import io
import html as _html
import streamlit as st
import streamlit.components.v1 as components
import pandas as pd

sys.path.insert(0, os.path.dirname(__file__))

try:
    from dotenv import load_dotenv
    load_dotenv(os.path.join(os.path.dirname(__file__), '.env'), override=False)
except ImportError:
    pass

from services.validator import validate_upload, ValidationError
from services.parser import parse_html, MissingFieldsError
from services.mapper import map_to_draft
from services.generator import generate_draft_text
from services.llm_fallback import extract_with_llm, LLMFallbackError
from models.schema import DraftDocument

st.set_page_config(
    page_title="SAP 입찰기안문 자동생성",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── 세션 상태 ──────────────────────────────────────────────────────────────────
_DEFAULTS = {
    "extracted_data": None,
    "draft_document": None,
    "warnings": [],
    "step": 0,
    "edited": {},
    "orig_fields": {},
    "manual_mode": False,
    "uploader_key": 0,
    "llm_confirm_pending": False,
    "llm_confirm_type": None,
    "llm_confirm_missing": [],
    "llm_confirmed": False,
}
for k, v in _DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ── 사이드바 ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.title("📋 SAP 입찰기안문")
    st.caption("자동생성 에이전트")
    st.divider()

    st.subheader("🔑 API 설정")
    api_key_input = st.text_input(
        "OpenAI API Key",
        value=os.environ.get("OPENAI_API_KEY", ""),
        type="password",
        placeholder="sk-...",
    )
    if api_key_input:
        os.environ["OPENAI_API_KEY"] = api_key_input
    else:
        os.environ.pop("OPENAI_API_KEY", None)

    model_input = st.selectbox(
        "모델",
        options=["gpt-3.5-turbo", "gpt-4", "gpt-4o-mini"],
        index=["gpt-3.5-turbo", "gpt-4", "gpt-4o-mini"].index(
            os.environ.get("OPENAI_MODEL", "gpt-3.5-turbo")
        ) if os.environ.get("OPENAI_MODEL", "gpt-3.5-turbo") in ["gpt-3.5-turbo", "gpt-4", "gpt-4o-mini"] else 0,
    )
    os.environ["OPENAI_MODEL"] = model_input

    st.divider()
    if st.button("🔄 초기화", use_container_width=True):
        _new_uploader_key = st.session_state.get("uploader_key", 0) + 1
        for k, v in _DEFAULTS.items():
            st.session_state[k] = v
        st.session_state.uploader_key = _new_uploader_key
        st.session_state.pop("current_file_key", None)
        st.rerun()


# ── 탭 분리 ───────────────────────────────────────────────────────────────────
tab1, tab2 = st.tabs(["📋 SAP 입찰기안문 (파일 업로드)", "📊 견적 비교 기안문 (직접 입력)"])


# ════════════════════════════════════════════════════════════════════════════════
# TAB 1 — SAP 입찰기안문 (기존 서비스)
# ════════════════════════════════════════════════════════════════════════════════
with tab1:
    # ── LLM 외부 전송 확인 팝업 ──────────────────────────────────────────────
    if st.session_state.get("llm_confirm_pending"):
        confirm_type = st.session_state.get("llm_confirm_type", "")

        if confirm_type == "pdf":
            reason = "PDF 파일은 규칙 기반 파싱이 불가능하여 **반드시 AI(OpenAI)를 통해 처리**됩니다."
        else:
            missing = st.session_state.get("llm_confirm_missing", [])
            reason = (
                f"규칙 기반 파싱으로 다음 필드를 추출하지 못했습니다: **{', '.join(missing)}**\n\n"
                "AI(OpenAI) 폴백으로 재시도하려고 합니다."
            )

        st.warning(
            f"⚠️ **외부 AI 처리 전 확인이 필요합니다**\n\n"
            f"{reason}\n\n"
            "**OpenAI 서버로 전송되는 정보:**\n"
            "- 입찰 파일의 텍스트 내용 (최대 8,000자)\n"
            "- 업체명 · 입찰금액 · 낙찰 정보 등 포함 가능\n\n"
            "사내 보안 정책에 따라 외부 전송이 제한되는 경우, **취소 후 수동 입력**을 이용해 주세요."
        )

        col_yes, col_no = st.columns(2)
        with col_yes:
            if st.button("✅ 외부 AI로 처리 진행", type="primary", use_container_width=True):
                st.session_state.llm_confirm_pending = False
                st.session_state.llm_confirmed = True
                st.rerun()
        with col_no:
            if st.button("❌ 취소 → 수동 입력으로 전환", use_container_width=True):
                st.session_state.llm_confirm_pending = False
                st.session_state.manual_mode = True
                st.rerun()
        st.stop()

    # ── 파일 업로드 ───────────────────────────────────────────────────────────
    st.header("파일 업로드")
    uploaded_file = st.file_uploader(
        "입찰 결과 파일을 드래그하거나 클릭하여 선택하세요.",
        type=["html", "pdf"],
        help="SAP 입찰 결과 HTML 또는 PDF 파일 (최대 10 MB)",
        key=f"file_uploader_{st.session_state.uploader_key}",
    )

    if uploaded_file:
        file_bytes = uploaded_file.read()
        file_key = f"{uploaded_file.name}_{len(file_bytes)}"

        if st.session_state.get("current_file_key") != file_key:
            _saved_uploader_key = st.session_state.get("uploader_key", 0)
            for k, v in _DEFAULTS.items():
                st.session_state[k] = v
            st.session_state.uploader_key = _saved_uploader_key
            st.session_state.current_file_key = file_key

        if st.session_state.step == 0:
            filename = uploaded_file.name
            file_ext = os.path.splitext(filename.lower())[1]

            extracted = None
            warnings = []

            if file_ext == ".pdf":
                if len(file_bytes) == 0:
                    st.error("❌ 빈 파일입니다.")
                    st.stop()
                if len(file_bytes) > 10 * 1024 * 1024:
                    st.error("❌ 파일 크기가 10MB를 초과합니다.")
                    st.stop()

                api_key = os.environ.get("OPENAI_API_KEY", "")
                if not api_key:
                    st.error("❌ PDF 파일 처리에는 OpenAI API 키가 필요합니다. 사이드바에서 API 키를 입력하세요.")
                    st.stop()

                try:
                    from services.pdf_extractor import extract_pdf_text
                    with st.spinner("PDF 텍스트 추출 중..."):
                        pdf_text = extract_pdf_text(file_bytes)
                except Exception as ex:
                    st.error(f"❌ PDF 읽기 실패: {ex}")
                    st.stop()

                if not st.session_state.get("llm_confirmed"):
                    st.session_state.llm_confirm_pending = True
                    st.session_state.llm_confirm_type = "pdf"
                    st.rerun()

                with st.spinner("AI로 데이터 추출 중..."):
                    try:
                        model_override = os.environ.get("OPENAI_MODEL", "")
                        extracted, warnings = extract_with_llm(pdf_text, api_key, model_override)
                        st.success(f"✅ PDF 데이터 추출 완료: {filename}")
                    except LLMFallbackError as le:
                        st.error(f"❌ PDF 처리 실패: {le}")
                        st.stop()
                st.session_state.llm_confirmed = False

            else:
                with st.spinner("파일 유효성 검사 중..."):
                    try:
                        html_content = validate_upload(filename, file_bytes)
                        st.success(f"✅ 파일 유효성 검사 통과: {filename}")
                    except ValidationError as e:
                        st.error(f"❌ [{e.code}] {e.message}")
                        st.stop()

                with st.spinner("HTML 파싱 및 데이터 추출 중..."):
                    try:
                        extracted, warnings = parse_html(html_content)
                    except MissingFieldsError as e:
                        api_key = os.environ.get("OPENAI_API_KEY", "")
                        if not api_key:
                            st.error("❌ API 키가 설정되지 않아 LLM 폴백을 사용할 수 없습니다. 수동 입력 모드로 전환합니다.")
                            st.session_state.manual_mode = True
                        elif not st.session_state.get("llm_confirmed"):
                            st.session_state.llm_confirm_pending = True
                            st.session_state.llm_confirm_type = "html_fallback"
                            st.session_state.llm_confirm_missing = e.missing_fields
                            st.rerun()
                        else:
                            try:
                                model_override = os.environ.get("OPENAI_MODEL", "")
                                extracted, warnings = extract_with_llm(html_content, api_key, model_override)
                                st.success("✅ LLM 폴백으로 데이터 추출 완료")
                            except LLMFallbackError as le:
                                st.error(f"❌ LLM 폴백도 실패했습니다: {le}\n수동 입력 모드로 전환합니다.")
                                st.session_state.manual_mode = True
                            st.session_state.llm_confirmed = False

            if st.session_state.manual_mode or extracted is None:
                st.subheader("✏️ 수동 입력")
                st.info("자동 추출에 실패했습니다. 아래에 직접 입력해 주세요.")

                with st.form("manual_form"):
                    rfx_name = st.text_input("RFx 명 / 입찰명")
                    selected_vendor = st.text_input("낙찰 업체명")
                    final_price_str = st.text_input("최종 입찰금액 (숫자만)")
                    payment_terms = st.text_input("대금지급조건")
                    vendor_list_str = st.text_area(
                        "참여 업체 목록 (업체명,금액 — 한 줄에 하나씩)",
                        placeholder="예:\n한솔테크(주),5000000\n대한상사,5500000",
                    )
                    submitted = st.form_submit_button("저장 후 계속", type="primary")

                if submitted:
                    from models.schema import ExtractedData, VendorBid
                    from services.cleaner import clean_rfx_name, clean_price, clean_vendor_name

                    vendors = []
                    for line in vendor_list_str.strip().split("\n"):
                        parts = line.split(",")
                        if len(parts) >= 2:
                            vendors.append(VendorBid(
                                name=clean_vendor_name(parts[0]),
                                unit_price=clean_price(parts[1]),
                            ))

                    extracted = ExtractedData(
                        rfx_name=rfx_name,
                        rfx_name_clean=clean_rfx_name(rfx_name),
                        vendors=vendors,
                        selected_vendor=clean_vendor_name(selected_vendor),
                        final_price=clean_price(final_price_str),
                        payment_terms=payment_terms,
                        extraction_method="manual",
                    )
                    warnings = ["수동 입력 모드로 데이터를 입력했습니다."]
                    st.session_state.manual_mode = False

            if extracted is not None:
                if not extracted.vendors:
                    st.error("❌ 업체 정보를 찾을 수 없습니다. 프로세스를 중단합니다.")
                    st.stop()

                draft, map_warnings = map_to_draft(extracted)
                draft.raw_text = generate_draft_text(draft)
                all_warnings = warnings + map_warnings

                st.session_state.extracted_data = extracted
                st.session_state.draft_document = draft
                st.session_state.warnings = all_warnings
                st.session_state.step = 2
                st.session_state.orig_fields = {
                    "order_vendor": draft.order_vendor,
                    "item_name": draft.item_name,
                    "line_items": "\n".join(draft.line_items),
                    "order_amount": draft.order_amount,
                    "payment_terms": draft.payment_terms,
                    "delivery_date": draft.delivery_date,
                    "purchase_reason": draft.purchase_reason,
                    "attachments": draft.attachments,
                }
                st.session_state.edited = dict(st.session_state.orig_fields)
                st.rerun()

    # ── 결과 표시 (step >= 2) ─────────────────────────────────────────────────
    if st.session_state.step >= 2:
        extracted = st.session_state.extracted_data
        draft = st.session_state.draft_document
        edited = st.session_state.edited
        orig = st.session_state.orig_fields

        for w in st.session_state.warnings:
            st.warning(f"⚠️ {w}")

        header_copy_placeholder = st.empty()
        code_placeholder = st.empty()

        st.divider()

        st.subheader("✏️ 기안문 편집")
        st.caption("항목을 수정하면 커서 이동 시 위 미리보기에 즉시 반영됩니다.")

        def field_label(key: str, label: str) -> str:
            return f"🟡 {label} *(수정됨)*" if edited.get(key) != orig.get(key) else label

        col1, col2 = st.columns(2)
        with col1:
            edited["order_vendor"] = st.text_input(
                field_label("order_vendor", "1. 발주업체"),
                value=edited.get("order_vendor", ""),
            )
            edited["line_items"] = st.text_area(
                field_label("line_items", "2. 구매품목 (항목별 한 줄씩)"),
                value=edited.get("line_items", ""),
                height=160,
                help="각 품목을 한 줄에 하나씩 입력하세요. ①②③ 번호는 자동으로 붙습니다.",
            )
            edited["payment_terms"] = st.text_input(
                field_label("payment_terms", "3. 대금지급조건"),
                value=edited.get("payment_terms", ""),
            )
        with col2:
            edited["delivery_date"] = st.text_input(
                field_label("delivery_date", "4. 납기일"),
                value=edited.get("delivery_date", ""),
                placeholder="예: 2025-12-31",
            )
            edited["purchase_reason"] = st.text_area(
                field_label("purchase_reason", "5. 구매요청사유"),
                value=edited.get("purchase_reason", ""),
                height=100,
            )
            edited["attachments"] = st.text_input(
                field_label("attachments", "7. 첨부자료"),
                value=edited.get("attachments", ""),
            )

        st.session_state.edited = edited

        st.subheader("업체 순위 목록")
        rank_df = pd.DataFrame(draft.vendor_rank_table)
        edited_rank_df = st.data_editor(
            rank_df,
            use_container_width=True,
            num_rows="dynamic",
            key="rank_table_editor",
        )

        _final_items = [l.strip() for l in edited.get("line_items", "").split("\n") if l.strip()]
        live_draft = DraftDocument(
            order_vendor=edited["order_vendor"],
            item_name=edited.get("item_name", ""),
            order_amount=edited.get("order_amount", ""),
            payment_terms=edited["payment_terms"],
            delivery_date=edited["delivery_date"],
            purchase_reason=edited["purchase_reason"],
            attachments=edited["attachments"],
            vendor_rank_table=edited_rank_df.to_dict("records"),
            line_items=_final_items,
        )
        live_draft.raw_text = generate_draft_text(live_draft)

        _esc_text = _html.escape(live_draft.raw_text)

        with header_copy_placeholder.container():
            components.html(f"""
            <textarea id="_ct" readonly style="position:fixed;top:-9999px;left:-9999px;font-size:1px;">{_esc_text}</textarea>
            <div style="display:flex;align-items:center;gap:10px;padding:6px 0 4px 0;">
              <span style="font-size:1.4rem;font-weight:700;
                font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;
                color:#31333F;">📄 기안문 미리보기</span>
              <button id="cpyBtn" onclick="(function(){{
                var btn = document.getElementById('cpyBtn');
                var ta = document.getElementById('_ct');
                var txt = ta.value;
                function ok(){{ btn.innerText='✅ 완료!'; setTimeout(function(){{btn.innerText='📋 복사';}},2000); }}
                function fb(){{ ta.removeAttribute('readonly'); ta.select(); try{{document.execCommand('copy');ok();}}catch(e){{btn.innerText='❌';}} ta.setAttribute('readonly',''); }}
                if(navigator.clipboard && window.isSecureContext){{navigator.clipboard.writeText(txt).then(ok,fb);}}else{{fb();}}
              }})();" style="padding:5px 14px;font-size:13px;font-weight:600;
                background:#1E88E5;color:white;border:none;border-radius:6px;
                cursor:pointer;white-space:nowrap;">📋 복사</button>
            </div>
            """, height=52)

        code_placeholder.code(live_draft.raw_text, language=None)

        st.divider()

        st.subheader("⬇️ 파일 내보내기")

        components.html(f"""
        <textarea id="_ct2" readonly style="position:fixed;top:-9999px;left:-9999px;font-size:1px;">{_esc_text}</textarea>
        <button id="copyAllBtn" onclick="(function(){{
          var btn = document.getElementById('copyAllBtn');
          var ta = document.getElementById('_ct2');
          var txt = ta.value;
          function ok(){{ btn.innerText='✅ 복사 완료!'; setTimeout(function(){{btn.innerText='📋 전체 기안문 한번에 복사';}},2500); }}
          function fb(){{ ta.removeAttribute('readonly'); ta.select(); try{{document.execCommand('copy');ok();}}catch(e){{btn.innerText='❌ 실패';}} ta.setAttribute('readonly',''); }}
          if(navigator.clipboard && window.isSecureContext){{navigator.clipboard.writeText(txt).then(ok,fb);}}else{{fb();}}
        }})();" style="
          width:100%; padding:14px 0; font-size:17px; font-weight:bold;
          background:#1E88E5; color:white; border:none; border-radius:8px;
          cursor:pointer; letter-spacing:0.5px; margin-bottom:8px;
        ">📋 전체 기안문 한번에 복사</button>
        """, height=62)

        exp_col1, exp_col2 = st.columns(2)

        txt_bytes = live_draft.raw_text.encode("utf-8")
        exp_col1.download_button(
            label="📄 .txt 다운로드",
            data=txt_bytes,
            file_name=f"SAP_기안문_{extracted.rfx_name_clean[:20]}.txt",
            mime="text/plain",
            use_container_width=True,
        )

        try:
            from docx import Document as DocxDocument

            def build_docx(draft_doc: DraftDocument) -> bytes:
                doc = DocxDocument()
                doc.add_heading("구매 기안문", 0)

                def add_row(table, label, value):
                    row = table.add_row()
                    row.cells[0].text = label
                    row.cells[1].text = value or ""

                info_table = doc.add_table(rows=0, cols=2)
                info_table.style = "Table Grid"
                for label, val in [
                    ("발주업체", draft_doc.order_vendor),
                    ("구매품목", draft_doc.item_name),
                    ("발주금액", draft_doc.order_amount),
                    ("대금지급조건", draft_doc.payment_terms),
                    ("납기일", draft_doc.delivery_date or ""),
                    ("구매요청사유", draft_doc.purchase_reason),
                    ("첨부자료", draft_doc.attachments),
                ]:
                    add_row(info_table, label, val)

                doc.add_heading("입찰업체 순위", level=1)
                rank_table = doc.add_table(rows=1, cols=3)
                rank_table.style = "Table Grid"
                hdr = rank_table.rows[0].cells
                hdr[0].text, hdr[1].text, hdr[2].text = "순위", "업체명", "입찰금액"
                for row in draft_doc.vendor_rank_table:
                    cells = rank_table.add_row().cells
                    cells[0].text = str(row.get("순위", ""))
                    cells[1].text = str(row.get("업체명", ""))
                    cells[2].text = str(row.get("입찰금액", ""))

                buf = io.BytesIO()
                doc.save(buf)
                return buf.getvalue()

            docx_bytes = build_docx(live_draft)
            exp_col2.download_button(
                label="📝 .docx 다운로드",
                data=docx_bytes,
                file_name=f"SAP_기안문_{extracted.rfx_name_clean[:20]}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except ImportError:
            exp_col2.warning("python-docx가 설치되지 않았습니다.")

        st.divider()

        with st.expander("🔍 추출 데이터 확인 (원본 ↔ 정제)", expanded=False):
            col_l, col_r = st.columns(2)
            with col_l:
                st.subheader("원본 추출")
                st.markdown(f"**RFx 명:** {extracted.rfx_name}")
                st.markdown(f"**선정 업체:** {extracted.selected_vendor}")
                st.markdown(f"**최종금액:** {extracted.final_price:,}")
                st.markdown(f"**대금지급조건:** {extracted.payment_terms or '(없음)'}")
                st.markdown(f"**추출 방법:** `{extracted.extraction_method}`")
                st.markdown("**참여 업체 목록:**")
                for v in extracted.vendors:
                    excl = " ~~제외~~" if v.is_excluded else ""
                    st.markdown(f"- {v.name}: {v.unit_price:,}원{excl}")
            with col_r:
                st.subheader("정제 후 매핑")
                st.markdown(f"**구매품목:** {draft.item_name}")
                st.markdown(f"**발주업체:** {draft.order_vendor}")
                st.markdown(f"**발주금액:** {draft.order_amount}")
                st.markdown(f"**대금지급조건:** {draft.payment_terms or '(미입력)'}")
                st.markdown("**업체 순위 목록:**")
                for row in draft.vendor_rank_table:
                    st.markdown(f"- {row['순위']}순위: {row['업체명']} / {row['입찰금액']}")


# ════════════════════════════════════════════════════════════════════════════════
# TAB 2 — 견적 비교 기안문 (직접 입력)
# ════════════════════════════════════════════════════════════════════════════════

def _parse_price(price_str: str) -> int:
    """'79.8만', '270만', '1.5억', '3,000,000' 등을 정수 원화로 변환."""
    s = price_str.strip().replace(",", "").replace(" ", "").replace("원", "")
    multiplier = 1
    if s.endswith("억"):
        multiplier = 100_000_000
        s = s[:-1]
    elif s.endswith("천만"):
        multiplier = 10_000_000
        s = s[:-2]
    elif s.endswith("백만"):
        multiplier = 1_000_000
        s = s[:-2]
    elif s.endswith("만"):
        multiplier = 10_000
        s = s[:-1]
    elif s.endswith("천"):
        multiplier = 1_000
        s = s[:-1]
    try:
        return int(round(float(s) * multiplier))
    except ValueError:
        raise ValueError(f"'{price_str}'을 금액으로 인식할 수 없습니다.")


def _parse_vendor_input(text: str) -> list:
    """'업체명 금액, 업체명 금액' 형식 파싱 후 금액 오름차순 정렬."""
    entries = [e.strip() for e in text.split(",") if e.strip()]
    if not entries:
        return []
    result = []
    for entry in entries:
        parts = entry.rsplit(None, 1)
        if len(parts) != 2:
            raise ValueError(f"'{entry}': 업체명과 금액을 공백으로 구분해 주세요.\n예) 중앙정밀 79.8만")
        name, price_str = parts
        amount = _parse_price(price_str)
        result.append({"name": name.strip(), "amount": amount})
    return sorted(result, key=lambda x: x["amount"])


def _fmt_amount(amount: int) -> str:
    return f"{amount:,}원"


def _generate_quote_draft(
    vendors: list,
    item_name: str,
    delivery_date: str,
    payment_terms: str,
) -> str:
    winner = vendors[0]
    all_names = [v["name"] for v in vendors]
    n = len(vendors)
    names_str = ", ".join(all_names)

    lines = [
        f"1. 발주업체: {winner['name']}",
        f"2. 구매품목: {item_name}",
        f"3. 발주금액: {_fmt_amount(winner['amount'])}",
        f"4. 대금지급조건 : {payment_terms}",
        f"5. 납기일: {delivery_date}",
        "6. 제안금액",
    ]
    for v in vendors:
        lines.append(f"   - {v['name']}: {_fmt_amount(v['amount'])}")
    lines.append("7. 구매요청사유")
    lines.append(
        f"전문업체 {n}개사({names_str}) 제안 후 최저가 업체인 {winner['name']}을 선정하였음."
    )
    lines.append("8. 첨부자료: 각 업체 제안서 1부씩")
    return "\n".join(lines)


with tab2:
    st.header("견적 비교 기안문")
    st.caption("업체명과 금액을 입력하면 최저가 업체 기준으로 기안문을 자동 생성합니다.")

    # 미리보기 영역 예약 (편집 위젯보다 시각적으로 위에 배치)
    q_header_placeholder = st.empty()
    q_code_placeholder = st.empty()

    st.divider()

    st.subheader("✏️ 입력")
    q_col1, q_col2 = st.columns(2)
    with q_col1:
        vendor_input = st.text_input(
            "업체명 금액 (쉼표로 구분)",
            placeholder="예: 중앙정밀 79.8만, 태광롤테크 270만",
            help="업체명과 금액을 공백으로 구분하고, 업체 간에는 쉼표(,)로 구분합니다.\n"
                 "금액 단위: 만 / 천만 / 백만 / 억 또는 숫자만 입력 (원 단위)",
        )
        q_item_name = st.text_input("구매품목", placeholder="(선택사항)")
    with q_col2:
        q_delivery = st.text_input("납기일", placeholder="예: 2025-12-31 (선택사항)")
        q_payment = st.text_input("대금지급조건", value="당사지급기준 만족")

    # 파싱 및 미리보기 생성
    q_vendors = []
    q_error = None
    if vendor_input.strip():
        try:
            q_vendors = _parse_vendor_input(vendor_input)
        except ValueError as e:
            q_error = str(e)

    if q_error:
        st.error(f"❌ {q_error}")

    if q_vendors:
        q_draft_text = _generate_quote_draft(q_vendors, q_item_name, q_delivery, q_payment)
        _q_esc = _html.escape(q_draft_text)

        with q_header_placeholder.container():
            components.html(f"""
            <textarea id="_qct" readonly style="position:fixed;top:-9999px;left:-9999px;font-size:1px;">{_q_esc}</textarea>
            <div style="display:flex;align-items:center;gap:10px;padding:6px 0 4px 0;">
              <span style="font-size:1.4rem;font-weight:700;
                font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;
                color:#31333F;">📄 기안문 미리보기</span>
              <button id="qCpyBtn" onclick="(function(){{
                var btn = document.getElementById('qCpyBtn');
                var ta = document.getElementById('_qct');
                var txt = ta.value;
                function ok(){{ btn.innerText='✅ 완료!'; setTimeout(function(){{btn.innerText='📋 복사';}},2000); }}
                function fb(){{ ta.removeAttribute('readonly'); ta.select(); try{{document.execCommand('copy');ok();}}catch(e){{btn.innerText='❌';}} ta.setAttribute('readonly',''); }}
                if(navigator.clipboard && window.isSecureContext){{navigator.clipboard.writeText(txt).then(ok,fb);}}else{{fb();}}
              }})();" style="padding:5px 14px;font-size:13px;font-weight:600;
                background:#1E88E5;color:white;border:none;border-radius:6px;
                cursor:pointer;white-space:nowrap;">📋 복사</button>
            </div>
            """, height=52)

        q_code_placeholder.code(q_draft_text, language=None)

        st.divider()

        components.html(f"""
        <textarea id="_qct2" readonly style="position:fixed;top:-9999px;left:-9999px;font-size:1px;">{_q_esc}</textarea>
        <button id="qCopyAllBtn" onclick="(function(){{
          var btn = document.getElementById('qCopyAllBtn');
          var ta = document.getElementById('_qct2');
          var txt = ta.value;
          function ok(){{ btn.innerText='✅ 복사 완료!'; setTimeout(function(){{btn.innerText='📋 전체 기안문 한번에 복사';}},2500); }}
          function fb(){{ ta.removeAttribute('readonly'); ta.select(); try{{document.execCommand('copy');ok();}}catch(e){{btn.innerText='❌ 실패';}} ta.setAttribute('readonly',''); }}
          if(navigator.clipboard && window.isSecureContext){{navigator.clipboard.writeText(txt).then(ok,fb);}}else{{fb();}}
        }})();" style="
          width:100%; padding:14px 0; font-size:17px; font-weight:bold;
          background:#1E88E5; color:white; border:none; border-radius:8px;
          cursor:pointer; letter-spacing:0.5px; margin-bottom:8px;
        ">📋 전체 기안문 한번에 복사</button>
        """, height=62)

        q_txt_bytes = q_draft_text.encode("utf-8")
        st.download_button(
            label="📄 .txt 다운로드",
            data=q_txt_bytes,
            file_name="견적비교_기안문.txt",
            mime="text/plain",
            use_container_width=True,
        )
    else:
        with q_header_placeholder.container():
            st.info("위에 업체명과 금액을 입력하면 기안문이 자동 생성됩니다.\n\n"
                    "예시: `중앙정밀 79.8만, 태광롤테크 270만`")
